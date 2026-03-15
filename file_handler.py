"""
file_handler.py
Read text from – and write translated text back to – various file formats.

Supported input formats : .txt  .docx  .pdf  .csv  .xlsx
Supported output formats: same as input (.pdf is saved as .txt)
"""

import csv
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf", ".csv", ".xlsx"}


class FileHandler:
    """Read content from and write content to various file formats."""

    # ── Reading (for preview) ─────────────────────────────────────────────────

    def read_file(self, filepath: str) -> str:
        """
        Extract all text from *filepath* and return it as a plain string.
        Used for the Original-tab preview only.
        Raises ValueError for unsupported extensions.
        """
        path = Path(filepath)
        ext = path.suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'.\n"
                f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        readers = {
            ".txt":  self._read_txt,
            ".docx": self._read_docx,
            ".pdf":  self._read_pdf,
            ".csv":  self._read_csv,
            ".xlsx": self._read_xlsx,
        }
        return readers[ext](path)

    def _read_txt(self, path: Path) -> str:
        for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                return path.read_text(encoding=enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return path.read_text(encoding="utf-8", errors="replace")

    def _read_docx(self, path: Path) -> str:
        doc = Document(str(path))
        lines: list[str] = []
        for para in doc.paragraphs:
            lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                lines.append("  |  ".join(row_cells))
        return "\n".join(lines)

    def _read_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n--- Page Break ---\n\n".join(pages)

    def _read_csv(self, path: Path) -> str:
        df = None
        for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                df = pd.read_csv(str(path), encoding=enc, header=None)
                break
            except UnicodeDecodeError:
                continue
        if df is None:
            df = pd.read_csv(str(path), encoding="latin-1", header=None)
        return self._dataframe_to_text(df)

    def _read_xlsx(self, path: Path) -> str:
        # Preview: only first sheet for speed
        df = pd.read_excel(str(path), engine="openpyxl", header=None, sheet_name=0)
        return self._dataframe_to_text(df)

    @staticmethod
    def _dataframe_to_text(df: "pd.DataFrame") -> str:
        lines: list[str] = []
        for _, row in df.iterrows():
            cells = [str(v).strip() for v in row
                     if str(v).strip() not in ("", "nan", "NaN")]
            if cells:
                lines.append("  |  ".join(cells))
        return "\n".join(lines)

    # ── Structure-preserving translation ─────────────────────────────────────

    def translate_file_structured(
        self,
        input_path: str,
        output_path: str,
        translate_fn,
        progress_callback=None,
    ) -> str:
        """
        Translate *input_path* while preserving its original file structure
        and save the result to *output_path*.

        *translate_fn*      – callable ``(text: str) -> str``
        *progress_callback* – callable ``(int 0-100)`` or None

        Returns a plain-text preview of the translated content.
        """
        ext = Path(input_path).suffix.lower()
        dispatch = {
            ".xlsx": self._translate_xlsx,
            ".csv":  self._translate_csv,
            ".docx": self._translate_docx_structured,
        }
        if ext in dispatch:
            return dispatch[ext](input_path, output_path, translate_fn, progress_callback)

        # txt / pdf – translate as plain text, save to output_path
        content    = self.read_file(input_path)
        translated = translate_fn(content)
        self._write_txt(Path(output_path), translated)
        return translated

    def _translate_xlsx(
        self,
        input_path: str,
        output_path: str,
        translate_fn,
        progress_callback=None,
    ) -> str:
        """Translate every text cell in every sheet; write a proper .xlsx file."""
        from openpyxl import load_workbook
        from openpyxl.cell.cell import MergedCell

        wb = load_workbook(str(input_path))
        preview_parts: list[str] = []

        # Collect all translatable cells across every sheet
        translatable: list = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    if (
                        not isinstance(cell, MergedCell)
                        and isinstance(cell.value, str)
                        and cell.value.strip()
                        and not cell.value.startswith("=")
                    ):
                        translatable.append(cell)

        total = max(len(translatable), 1)
        cache: dict[str, str] = {}          # avoid re-translating identical text
        seen_sheets: set[str] = set()

        for i, cell in enumerate(translatable):
            sheet_title = cell.parent.title
            if sheet_title not in seen_sheets:
                preview_parts.append(f"\n=== {sheet_title} ===")
                seen_sheets.add(sheet_title)

            original = cell.value
            if original not in cache:
                cache[original] = translate_fn(original)
            cell.value = cache[original]
            preview_parts.append(cell.value)

            if progress_callback:
                progress_callback(int((i + 1) / total * 100))

        wb.save(str(output_path))
        return "\n".join(preview_parts).strip()

    def _translate_csv(
        self,
        input_path: str,
        output_path: str,
        translate_fn,
        progress_callback=None,
    ) -> str:
        """Translate every non-empty cell; write a proper .csv file."""
        rows: list[list[str]] = []
        for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                with open(str(input_path), "r", encoding=enc, newline="") as f:
                    rows = list(csv.reader(f))
                break
            except UnicodeDecodeError:
                continue

        # Index all translatable (row, col, value) entries
        translatable = [
            (r, c, val)
            for r, row in enumerate(rows)
            for c, val in enumerate(row)
            if val.strip()
        ]

        total = max(len(translatable), 1)
        cache: dict[str, str] = {}
        preview_parts: list[str] = []

        for i, (r, c, val) in enumerate(translatable):
            if val not in cache:
                cache[val] = translate_fn(val)
            rows[r][c] = cache[val]
            preview_parts.append(cache[val])
            if progress_callback:
                progress_callback(int((i + 1) / total * 100))

        with open(str(output_path), "w", encoding="utf-8", newline="") as f:
            csv.writer(f).writerows(rows)

        return "\n".join(preview_parts)

    def _translate_docx_structured(
        self,
        input_path: str,
        output_path: str,
        translate_fn,
        progress_callback=None,
    ) -> str:
        """Translate every paragraph (body + tables); write a proper .docx file."""
        doc = Document(str(input_path))

        # Collect paragraphs that contain text (body and table cells)
        paras: list = []
        for p in doc.paragraphs:
            if p.text.strip():
                paras.append(p)

        seen_cell_ids: set[int] = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cid = id(cell)
                    if cid in seen_cell_ids:
                        continue
                    seen_cell_ids.add(cid)
                    for p in cell.paragraphs:
                        if p.text.strip():
                            paras.append(p)

        total = max(len(paras), 1)
        cache: dict[str, str] = {}
        preview_parts: list[str] = []

        for i, para in enumerate(paras):
            original = para.text
            if original not in cache:
                cache[original] = translate_fn(original)
            translated = cache[original]

            # Replace text while keeping first run's character formatting
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = translated
            else:
                para.add_run(translated)

            preview_parts.append(translated)
            if progress_callback:
                progress_callback(int((i + 1) / total * 100))

        doc.save(str(output_path))
        return "\n".join(preview_parts)

    # ── Plain-text writing (txt / pdf output) ─────────────────────────────────

    def write_file(
        self, output_path: str, content: str, original_path: str
    ) -> None:
        """Write plain-text *content* to *output_path* (used for txt/pdf)."""
        path     = Path(output_path)
        orig_ext = Path(original_path).suffix.lower()
        if orig_ext == ".docx":
            self._write_docx(path, content)
        else:
            self._write_txt(path, content)

    def _write_txt(self, path: Path, content: str) -> None:
        path.write_text(content, encoding="utf-8")

    def _write_docx(self, path: Path, content: str) -> None:
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(path))

    # ── Helpers ───────────────────────────────────────────────────────────────

    def get_output_extension(self, original_path: str) -> str:
        """Return the extension used when saving a translation.
        PDF is always saved as .txt (recreating PDF layout is impractical).
        """
        ext = Path(original_path).suffix.lower()
        return ".txt" if ext == ".pdf" else ext
